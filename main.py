import os
import uvicorn
import numpy as np
from typing import List
from docx import Document
from sentence_transformers import SentenceTransformer
from mcp.server.fastmcp import FastMCP
from dotenv import load_dotenv
from docx.shared import Pt

print(f"--- DEBUG: start of execution ---")
# Configuration constants
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
SIMILARITY_THRESHOLD = 0.3
TOP_K_RESULTS = 2


# This loads the variables from your .env file into os.environ
load_dotenv()
# Global state
mcp = FastMCP("Sentinel-Knowledge-Base")

# Initialize SentenceTransformer model (lazy loading - only when needed)
model: SentenceTransformer = None

# Global storage for chunks and embeddings
chunks: List[str] = []
embeddings: np.ndarray = np.array([])

KNOWLEDGE_BASE_PATH = "knowledge_base.docx"


def get_model() -> SentenceTransformer:
    """
    Lazy load the SentenceTransformer model.
    Only loads when first needed to save startup time and resources.
    """
    global model
    if model is None:
        print("Loading SentenceTransformer model...")
        model = SentenceTransformer('all-MiniLM-L6-v2')
        print("Model loaded successfully")
    return model


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    Split text into overlapping chunks using a sliding window approach.

    Args:
        text: The input text to chunk
        size: Maximum size of each chunk in characters
        overlap: Number of overlapping characters between chunks

    Returns:
        List of text chunks
    """
    if not text or len(text) <= size:
        return [text] if text else []

    chunks = []
    start = 0

    while start < len(text):
        end = start + size

        # Try to find a good break point (sentence boundary)
        if end < len(text):
            # Look for sentence-ending punctuation within the last 50 characters
            break_candidates = [
                text.rfind('.', start, end),
                text.rfind('!', start, end),
                text.rfind('?', start, end),
                text.rfind('\n', start, end),
            ]

            # Use the latest valid break point that's within reasonable range
            best_break = max([b for b in break_candidates if b > start + size // 2], default=end - 1)
            end = best_break + 1

        chunks.append(text[start:end].strip())
        start = end - overlap

    return [c for c in chunks if c]


def compute_similarities(query_embedding: np.ndarray) -> np.ndarray:
    """
    Compute cosine similarities between query and all stored embeddings using NumPy.

    Args:
        query_embedding: Query vector (shape: [embedding_dim])

    Returns:
        Array of similarity scores (shape: [num_chunks])
    """
    global embeddings

    if embeddings.size == 0:
        return np.array([])

    # Reshape query to (1, embedding_dim) for matrix operations
    query = query_embedding.reshape(1, -1)

    # Calculate dot product (query embeddings @ stored embeddings.T)
    dot_product = np.dot(query, embeddings.T).flatten()

    # Calculate L2 norms
    query_norm = np.linalg.norm(query)
    embeddings_norm = np.linalg.norm(embeddings, axis=1)

    # Avoid division by zero
    if query_norm == 0 or np.any(embeddings_norm == 0):
        return np.zeros(len(embeddings))

    # Cosine similarity = (A · B) / (||A|| * ||B||)
    similarities = dot_product / (query_norm * embeddings_norm)

    return similarities


def index_document() -> None:
    """
    Index the knowledge base document by reading, chunking, and embedding.

    Reads 'knowledge_base.docx', applies sliding window chunking,
    generates embeddings using SentenceTransformer, and stores them in global state.
    Handles cases where the document doesn't exist or is empty.
    """
    global chunks, embeddings

    doc_path = "knowledge_base.docx"

    # Handle case where document doesn't exist
    if not os.path.exists(doc_path):
        print(f"Knowledge base document not found at {doc_path}")
        print("Create the document first using the /admin/feed endpoint")
        chunks = []
        embeddings = np.array([])
        return

    try:
        # Read the Word document
        doc = Document(doc_path)

        # Extract paragraphs, ignoring empty ones
        full_text = "\n".join([
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        ])

        # Handle empty document
        if not full_text:
            print("Knowledge base document is empty")
            chunks = []
            embeddings = np.array([])
            return

        # Apply sliding window chunking
        new_chunks = chunk_text(full_text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP)

        # Generate embeddings for all chunks
        new_embeddings = get_model().encode(
            new_chunks,
            convert_to_numpy=True,
            show_progress_bar=True
        )

        # Update global state
        chunks = new_chunks
        embeddings = new_embeddings

        print(f"Successfully indexed {len(chunks)} chunks from {doc_path}")
        print(f"Chunk size: {CHUNK_SIZE}, Overlap: {CHUNK_OVERLAP}")

    except Exception as e:
        print(f"Error indexing document: {e}")
        chunks = []
        embeddings = np.array([])


@mcp.tool(
    name="search_knowledge_base",
    description="Searches the internal Word documents for knowledge base whenever user asks about some information."
)
def search_knowledge_base(query: str) -> str:
    """
    Search the knowledge base for relevant chunks based on a query.

    Uses cosine similarity to find the most relevant chunks.
    Returns a friendly message if no good matches are found.

    Args:
        query: The search query string

    Returns:
        A formatted string containing the most relevant text chunks,
        or a friendly message if no good matches are found.
    """
    global chunks, embeddings

    # Handle empty knowledge base
    if len(chunks) == 0 or embeddings.size == 0:
        return "I don't have any knowledge base data yet. Please add content first."

    # Generate embedding for the query
    query_embedding = get_model().encode([query], convert_to_numpy=True)[0]

    # Compute similarities using NumPy-based cosine similarity
    similarities = compute_similarities(query_embedding)

    # Get indices of top k most similar chunks
    top_k = min(TOP_K_RESULTS, len(similarities))
    top_indices = np.argsort(similarities)[::-1][:top_k]

    # Filter results by similarity threshold
    valid_results = []
    for rank, idx in enumerate(top_indices, 1):
        similarity_score = similarities[idx]

        # Skip results below threshold
        if similarity_score < SIMILARITY_THRESHOLD:
            break

        chunk_text = chunks[idx]
        valid_results.append({
            "rank": rank,
            "similarity": similarity_score,
            "text": chunk_text
        })

    # Return friendly message if no good matches found
    if not valid_results:
        return "I don't know this yet. Try rephrasing your question or add more relevant information to the knowledge base."

    # Format results
    results = []
    for result in valid_results:
        results.append(
            f"[Result {result['rank']}] (similarity: {result['similarity']:.4f}):\n{result['text']}"
        )

    return "\n\n".join(results)

@mcp.tool(
    name="append_to_knowledge_base",
    description="Append new knowledge records or information to the Word document. Use this to 'feed' the bot new facts to remember"
)
async def append_to_knowledge_base(text: str) -> str:
    """
    Append new knowledge records or information to the Word document.
    Use this to 'feed' the bot new facts to remember.
    """
    try:
        # 1. Create or Load the document
        if os.path.exists(KNOWLEDGE_BASE_PATH):
            doc = Document(KNOWLEDGE_BASE_PATH)
        else:
            doc = Document()
            title = doc.add_heading("Knowledge Base", 0)
            title.runs[0].font.size = Pt(14)

        # 2. Append the text
        doc.add_paragraph(text)
        doc.save(KNOWLEDGE_BASE_PATH)

        # 3. 🚀 CRITICAL: Re-run your indexing logic here
        # This ensures the new text is added to your search embeddings immediately
        index_document() 

        return f"✅ Successfully added to knowledge base: {text[:50]}..."
        
    except Exception as e:
        return f"❌ Error updating knowledge base: {str(e)}"




if __name__ == "__main__":
    try:
        print(f"--- inside main ---")

        # Use PORT from environment (Render sets this), default to 8000
        port = int(os.environ.get("PORT", 8000))
        host = "0.0.0.0"  # Bind to all interfaces for Render

        print(f"Starting server on {host}:{port}")

        # Get the SSE app and run with uvicorn
        app = mcp.sse_app()
        uvicorn.run(app, host=host, port=port)

    except Exception as e:
        print(f"--- CRITICAL ERROR DURING STARTUP ---")
        print(str(e))
      